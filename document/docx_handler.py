import os
from docx import Document
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt
from pydocx import PyDocX
import re
import json
from document.translate import translate_en_to_cn, translate_type
import base64
import datetime
from PIL import Image
from ManagementSystemPremium.settings import BASE_DIR
from ManagementSystemPremium import safe
from urllib import parse


source_dir = os.path.join(BASE_DIR, "document/static/document/docs/source/")
templates_dir = os.path.join(BASE_DIR, "document/static/document/docs/templates/")
storage_dir = os.path.join(BASE_DIR, "document/storage/")
temporary_dir = os.path.join(BASE_DIR, "document/temporary/")
translate_dir = os.path.join(BASE_DIR, "document/static/document/docs/translate/")
translate_file_name = "translate.json"
translate_path = os.path.join(translate_dir, translate_file_name)


def check_dir_exists(dir_path):
    if not os.path.exists(dir_path):
        os.mkdir(dir_path)


for dir_path in [source_dir, templates_dir, storage_dir, temporary_dir, translate_dir]:
    check_dir_exists(dir_path)


def to_list(origin_text, is_type, is_not_type):
    if type(origin_text) is is_type:
        return [origin_text]
    elif type(origin_text) is not is_not_type:
        return None


def get_file_path(file_dir, file_name):
    file_path = os.path.join(file_dir, file_name)
    if os.path.isfile(file_path) and file_path.endswith(".docx"):
        return file_path
    else:
        return None


def create_docx_handler(docx_path, need_handler=''):
    if need_handler == 'python-docx':
        document_handler = Document(docx_path)
        document_template_handler = None
    elif need_handler == 'python-docx-template':
        document_handler = None
        document_template_handler = DocxTemplate(docx_path)
    else:
        document_handler = Document(docx_path)
        document_template_handler = DocxTemplate(docx_path)
    return document_handler, document_template_handler


def get_variable_list(document_handler, define_text='' or [], define_column_index=0 or [] or -1, search_regular=r".*_[0-9]+$"):
    table = document_handler.tables[0]
    variable_dict = {}
    target_list = []
    define_text = to_list(define_text, str, list)
    define_column_index = to_list(define_column_index, int, list) if define_column_index != -1 else list(range(len(table.columns)))
    search_regular = re.compile(search_regular)
    if define_text is None or define_column_index is None:
        return None
    for d_text in define_text:
        for row in range(len(table.rows)):
            for column in define_column_index:
                try:
                    result = table.cell(row, column).text
                except:
                    continue
                if d_text in result:
                    targets = search_regular.findall(result)
                    if len(targets) > 0:
                        if targets not in target_list:
                            target_list.append(targets)
                            for target in targets:
                                variable_slice = target.split("_")
                                try:
                                    variable_dict[variable_slice[0]] = {
                                        "en_name": variable_slice[0],
                                        "cn_name": translate_en_to_cn(variable_slice[0], translate_path),
                                        "type": variable_slice[1],
                                        "type_text": translate_type(variable_slice[1][1]),
                                        "maximum": variable_slice[2],
                                        "origin": target,
                                    }
                                except:
                                    variable_dict[variable_slice[0]] = {}
    return variable_dict


def docx_to_html(docx_path):
    docx_html = PyDocX.to_html(docx_path)
    style_start = docx_html.find("<style>")
    style_end = docx_html.find("</style>")
    docx_html = docx_html[:style_start] + docx_html[style_end:]
    docx_html = '<div>' + docx_html.replace("_", "&ensp;") + '</div>'
    return docx_html


def fix_image_rotate(image_path):
    img = Image.open(image_path)
    x, y = img.size
    if y > x:
        img = img.rotate(90, expand=True)
        img.save(image_path)
    img.close()


def write(document_template_handler, save_path, init_content: str = None, docx_content: list = None,
          docx_variable_dict: dict = None, auto_variable_dict: str = None, signature_content: list = None,
          supervisor_variable_dict: dict = None, supervisor_content_list: list = None, maximum=0, show_signature=False,
          show_signature_id=False):
    context = {}
    image_path_list = []
    i = 0
    if init_content:
        try:
            context.update(json.loads(init_content))
        except:
            pass
    if docx_content:
        for content in docx_content:
            try:
                content_dict = json.loads(content["content"])
                for key, value in content_dict.items():
                    if i == maximum-1:
                        name = docx_variable_dict[key]["origin"]
                    else:
                        name = key + '__' + str(i)
                    context.update({name: value})
            except:
                pass
            try:
                signature_name, signature_img, content_str = '', '', ''
                signature_id = content["signature_id"]
                if i == maximum-1:
                    name = docx_variable_dict["signature"]["origin"]
                else:
                    name = "signature" + '__' + str(i)
                if signature_id != "":
                    for s in signature_content:
                        if s['signature_id'] == signature_id:
                            signature_name = s['user__full_name']
                            signature_img = s['signature']
                    if show_signature:
                        if signature_img != '':
                            k, s = signature_img.split(';')
                            signature_img = base64.b64decode(parse.unquote(safe.aes_decrypt(s, k)).split(',')[1])
                            image_path = temporary_dir + signature_id + '.png'
                            with open(image_path, 'wb') as f:
                                f.write(signature_img)
                            fix_image_rotate(image_path)
                            content_str = InlineImage(document_template_handler, image_path, height=Mm(9))
                            image_path_list.append(image_path)
                    elif show_signature_id:
                        content_str = "{{ " + signature_id + " }}"
                    else:
                        if signature_name != '':
                            content_str = signature_name + ' 已签名'
                context.update({name: content_str})
            except:
                pass
            i += 1
    if auto_variable_dict:
        try:
            context.update(json.loads(auto_variable_dict))
        except:
            pass
    if supervisor_content_list:
        for content in supervisor_content_list:
            try:
                name = supervisor_variable_dict[content["content_name"]]["origin"]
                if content["content"]:
                    context.update({name: content["content"]})
                elif content["signature_id"]:
                    signature_name, signature_img, content_str = '', '', ''
                    signature_id = content["signature_id"]
                    if signature_id != "":
                        for s in signature_content:
                            if s['signature_id'] == signature_id:
                                signature_name = s['user__full_name']
                                signature_img = s['signature']
                        if show_signature:
                            if signature_img != '':
                                k, s = signature_img.split(';')
                                signature_img = base64.b64decode(parse.unquote(safe.aes_decrypt(s, k)).split(',')[1])
                                image_path = temporary_dir + signature_id + '.png'
                                with open(image_path, 'wb') as f:
                                    f.write(signature_img)
                                fix_image_rotate(image_path)
                                content_str = InlineImage(document_template_handler, image_path, height=Mm(9))
                                image_path_list.append(image_path)
                        elif show_signature_id:
                            content_str = "{{ " + signature_id + " }}"
                        else:
                            if signature_name != '':
                                content_str = signature_name + ' 已签名'
                    context.update({name: content_str})
            except:
                pass
            i += 1
    document_template_handler.render(context)
    document_template_handler.save(save_path)
    for image_path in image_path_list:
        os.remove(image_path)
